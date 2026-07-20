"""Tests for the new file-type parsers: CSV (stdlib), .xlsx (openpyxl),
.xls (xlrd), and .doc (antiword, exercised via its not-installed failure
path in this dev environment -- the success path is covered on the CI/prod
image where antiword is actually present). Multi-sheet coverage matters:
Excel workbooks with multiple tabs must all be captured, not just the first.
"""

from io import BytesIO

import pytest

from app.parser import CsvParser, DocParser, DocxParser, ExcelParser, XlsParser, parse_document


class TestCsvParser:
    async def test_parses_rows_into_text(self):
        content = b"name,role\nAlice,admin\nBob,viewer\n"
        result = await CsvParser.parse(content)
        assert result.status == "success"
        assert "Alice" in result.raw_text
        assert "admin" in result.raw_text

    async def test_empty_csv_is_partial(self):
        result = await CsvParser.parse(b"")
        assert result.status == "partial"


class TestExcelParser:
    def _workbook_bytes(self, sheet_data: dict[str, list[list]]) -> bytes:
        from openpyxl import Workbook

        wb = Workbook()
        wb.remove(wb.active)
        for name, rows in sheet_data.items():
            ws = wb.create_sheet(title=name)
            for row in rows:
                ws.append(row)
        buf = BytesIO()
        wb.save(buf)
        return buf.getvalue()

    async def test_single_sheet(self):
        content = self._workbook_bytes({"Sheet1": [["a", "b"], [1, 2]]})
        result = await ExcelParser.parse(content)
        assert result.status == "success"
        assert "Sheet1" in result.raw_text

    async def test_multiple_sheets_all_captured(self):
        """Every tab must show up in raw_text/sections, not just the first."""
        content = self._workbook_bytes(
            {
                "Summary": [["total", 100]],
                "Details": [["line item", "amount"], ["widget", 50]],
                "Notes": [["remember to review section 4"]],
            }
        )
        result = await ExcelParser.parse(content)
        assert result.status == "success"
        assert {s.heading for s in result.sections} == {"Summary", "Details", "Notes"}
        assert "widget" in result.raw_text
        assert "review section 4" in result.raw_text
        assert result.page_count == 3


class TestXlsParser:
    async def test_invalid_bytes_fails_gracefully(self):
        result = await XlsParser.parse(b"not a real xls file")
        assert result.status == "failed"
        assert result.error_message


class TestDocxParser:
    """A real-world regression: many SOW/RFP templates (this project's
    docs/sample/SOW_Template/*.docx among them) lay their content out in
    tables rather than plain paragraphs. The old paragraphs-only extraction
    silently returned status="success" with raw_text="" for every one of
    them -- a document with real content parsing to empty text, with no
    error surfaced anywhere."""

    def _docx_bytes(self, build_fn) -> bytes:
        from docx import Document

        doc = Document()
        build_fn(doc)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def test_table_only_document_extracts_text(self):
        """A docx with ALL its content in a table (no body paragraphs) must
        not come back empty -- this was the exact shape of the bug."""

        def build(doc):
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Project Title:"
            table.cell(0, 1).text = "IT Modernization"
            table.cell(1, 0).text = "Liability Cap:"
            table.cell(1, 1).text = "12 months of fees"

        result = await DocxParser.parse(self._docx_bytes(build))

        assert result.status == "success"
        assert "IT Modernization" in result.raw_text
        assert "Liability Cap" in result.raw_text
        assert "12 months of fees" in result.raw_text

    async def test_merged_cells_are_not_duplicated(self):
        """A merged cell spans multiple grid positions in python-docx but
        is the SAME underlying cell -- must appear once per row, not once
        per spanned column."""

        def build(doc):
            table = doc.add_table(rows=1, cols=3)
            table.cell(0, 0).merge(table.cell(0, 2))
            table.cell(0, 0).text = "SCOPE OF WORK"

        result = await DocxParser.parse(self._docx_bytes(build))

        assert result.status == "success"
        assert result.raw_text.count("SCOPE OF WORK") == 1

    async def test_paragraphs_and_tables_both_captured(self):
        """Headings/paragraphs still work exactly as before; tables are
        additive, not a replacement extraction path."""

        def build(doc):
            doc.add_heading("Statement of Work", level=1)
            doc.add_paragraph("This SOW covers the following deliverables.")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Deliverable"
            table.cell(0, 1).text = "Responsive website"

        result = await DocxParser.parse(self._docx_bytes(build))

        assert result.status == "success"
        assert "Statement of Work" in result.raw_text
        assert "following deliverables" in result.raw_text
        assert "Responsive website" in result.raw_text
        assert result.sections[0].heading == "Statement of Work"
        assert "Responsive website" in result.sections[0].content


class TestDocParser:
    async def test_missing_antiword_fails_gracefully(self):
        """In this dev environment antiword isn't installed -- confirms the
        graceful-failure path rather than an unhandled crash. The success
        path (antiword actually extracting text) runs on the prod image,
        which installs antiword via Dockerfile.prod."""
        result = await DocParser.parse(b"fake doc bytes")
        assert result.status == "failed"
        assert result.error_message


class TestParseDocumentDispatch:
    async def test_dispatches_csv_by_extension_alias(self):
        result = await parse_document(b"a,b\n1,2\n", "csv")
        assert result.status == "success"

    async def test_dispatches_xlsx_by_mime_type(self):
        from openpyxl import Workbook

        wb = Workbook()
        wb.active.append(["x", "y"])
        buf = BytesIO()
        wb.save(buf)

        result = await parse_document(
            buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        assert result.status == "success"

    async def test_unsupported_type_fails(self):
        result = await parse_document(b"data", "application/zip")
        assert result.status == "failed"
